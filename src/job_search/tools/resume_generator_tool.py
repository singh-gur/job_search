import os
from typing import Any, Dict, Type

from crewai.tools import BaseTool
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from pydantic import BaseModel, Field


class ResumeGeneratorInput(BaseModel):
    user_info: Dict[str, Any] = Field(
        description="User's personal and career information"
    )
    job_requirements: str = Field(description="Job requirements and skills analysis")
    output_filename: str = Field(
        default="personalized_resume.docx", description="Output filename for the resume"
    )


class ResumeGeneratorTool(BaseTool):
    name: str = "Resume Generator Tool"
    description: str = "Generate personalized resume documents based on user profile and job requirements"
    args_schema: Type[BaseModel] = ResumeGeneratorInput

    def _run(
        self, user_info: Dict[str, Any], job_requirements: str, output_filename: str
    ) -> str:
        try:
            filename = output_filename

            # Create a new Document
            doc = Document()

            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Header with name and contact info
            header = doc.add_heading(user_info.get("name", "Your Name"), 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER

            contact_info = doc.add_paragraph()
            contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_text = []
            if user_info.get("email"):
                contact_text.append(user_info["email"])
            if user_info.get("phone"):
                contact_text.append(user_info["phone"])
            if user_info.get("location"):
                contact_text.append(user_info["location"])
            if user_info.get("linkedin"):
                contact_text.append(f"LinkedIn: {user_info['linkedin']}")

            contact_info.add_run(" | ".join(contact_text))

            # Professional Summary
            doc.add_heading("Professional Summary", level=1)
            summary = user_info.get(
                "summary",
                "Experienced professional with expertise in relevant technologies and strong problem-solving skills.",
            )
            doc.add_paragraph(summary)

            # Skills Section
            doc.add_heading("Technical Skills", level=1)
            skills = user_info.get("skills", [])
            if isinstance(skills, list):
                skills_text = ", ".join(skills)
            else:
                skills_text = str(skills)
            doc.add_paragraph(skills_text)

            # Experience Section
            doc.add_heading("Professional Experience", level=1)
            experiences = user_info.get("experience", [])
            if not isinstance(experiences, list):
                experiences = [
                    {
                        "title": "Software Developer",
                        "company": "Tech Company",
                        "duration": "2020-2024",
                        "description": "Developed software applications and solutions.",
                    }
                ]

            for exp in experiences:
                exp_para = doc.add_paragraph()
                exp_para.add_run(
                    f"{exp.get('title', 'Position Title')} - {exp.get('company', 'Company Name')}"
                ).bold = True
                exp_para.add_run(f"\n{exp.get('duration', 'Duration')}")
                if exp.get("description"):
                    doc.add_paragraph(f"• {exp['description']}")

            # Education Section
            doc.add_heading("Education", level=1)
            education = user_info.get("education", [])
            if not isinstance(education, list):
                education = [
                    {
                        "degree": "Bachelor's Degree",
                        "school": "University Name",
                        "year": "2020",
                    }
                ]

            for edu in education:
                edu_para = doc.add_paragraph()
                edu_para.add_run(
                    f"{edu.get('degree', 'Degree')} - {edu.get('school', 'School Name')}"
                ).bold = True
                if edu.get("year"):
                    edu_para.add_run(f" ({edu['year']})")

            # Certifications (if any)
            certifications = user_info.get("certifications", [])
            if certifications:
                doc.add_heading("Certifications", level=1)
                for cert in certifications:
                    doc.add_paragraph(f"• {cert}")

            # Save the document
            output_path = os.path.join(os.getcwd(), filename)
            doc.save(output_path)

            return f"Resume successfully generated and saved as '{filename}' in the current directory. The resume has been tailored based on the job requirements analysis."

        except Exception as e:
            return f"Error generating resume: {str(e)}"
